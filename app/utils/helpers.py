import os
import re
from pathlib import Path
from typing import List, Tuple, Union
from pypdf import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import unicodedata
from datetime import datetime
import hashlib
from fastapi import HTTPException


def read_document(file_path: Union[str, Path]) -> str:
    """Read either PDF or DOCX file from compliance/contracts folders."""
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    if file_path.suffix.lower() == '.pdf':
        return _read_pdf(file_path)
    elif file_path.suffix.lower() == '.docx':
        return _read_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")


def _read_pdf(file_path: Path) -> str:
    """Internal PDF reader with legal document optimizations"""
    try:
        reader = PdfReader(file_path)
        return "\n\n".join(
            page.extract_text().strip()
            for page in reader.pages
            if page.extract_text().strip()
        )
    except Exception as e:
        raise ValueError(f"PDF read error: {file_path.name} - {str(e)}")


def _read_docx(file_path: Path) -> str:
    """Internal DOCX reader preserving legal formatting"""
    try:
        doc = Document(file_path)
        return "\n\n".join(
            para.text.strip()
            for para in doc.paragraphs
            if para.text.strip()
        )
    except Exception as e:
        raise ValueError(f"DOCX read error: {file_path.name} - {str(e)}")


def chunk_text(text: str, doc_type: str, method: str = "default") -> List[Tuple[str, dict]]:
    """
    Enhanced chunking with multiple methods.
    Options for method: "default", "legal_headers"
    """
    text = clean_text(text)

    if method == "legal_headers":
        sections = legal_headers_splitter(text)
        return [
            (section, {"type": doc_type, "chunk_method": "legal_headers"})
            for section in sections if section.strip()
        ]
    else:
        # Default recursive splitting
        if doc_type.lower() == 'contracts':
            chunk_size = 800
            overlap = 150
        else:
            chunk_size = 1000
            overlap = 200

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            separators=["\n\nSECTION", "\n\nArticle", "\n\nClause", "\n\n", "\n", " "],
            length_function=len
        )
        return [
            (chunk.page_content, {
                "type": doc_type,
                "chunk_method": "recursive",
                "char_count": len(chunk.page_content)
            })
            for chunk in splitter.create_documents([text])
        ]


def get_documents_from_folder(folder: str) -> List[Tuple[str, str]]:
    """
    Scan contract/ or compliance/ folder and return (file_path, doc_type) pairs.
    Doc_type is either 'contract' or the compliance standard name (e.g. 'GDPR').
    """
    base_path = Path("data") / folder
    if not base_path.exists():
        return []

    documents = []
    for file in base_path.glob("*"):
        if file.suffix.lower() in ('.pdf', '.docx'):
            doc_type = folder if folder == "contracts" else file.stem
            documents.append((str(file), doc_type))

    return documents

def clean_text(text: str) -> str:
    """
    Normalize legal text by:
    - Standardizing quotes/dashes
    - Removing non-printable chars
    - Normalizing unicode
    """
    # Normalize unicode first
    text = unicodedata.normalize('NFKC', text)

    # Replace legal document artifacts
    replacements = {
        '“': '"', '”': '"', '‘': "'", '’': "'",
        '":': '"', '":': '"', '': "'", '': "'",
        '–': '-', '—': '-', '§': 'Section',
        '\xa0': ' ', '\xad': '', '\u202f': ' '
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def generate_document_id(file_path: Union[str, Path]) -> str:
    """Create a unique ID based on file contents"""
    file_path = Path(file_path)
    hasher = hashlib.sha256()
    with file_path.open('rb') as f:
        while chunk := f.read(8192):
            hasher.update(chunk)
    return hasher.hexdigest()[:16]

def extract_metadata(file_path: Union[str, Path]) -> dict:
    """Extract basic document metadata"""
    path = Path(file_path)
    return {
        "filename": path.name,
        "extension": path.suffix.lower(),
        "size": path.stat().st_size,
        "modified": datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
        "created": datetime.fromtimestamp(path.stat().st_ctime).isoformat()
    }

def validate_document(file_path: Union[str, Path]) -> bool:
    """Verify document is valid and readable"""
    try:
        path = Path(file_path)
        if path.suffix.lower() == '.pdf':
            PdfReader(path).pages[0].extract_text()
        elif path.suffix.lower() == '.docx':
            Document(path).paragraphs[0].text
        return True
    except:
        return False

def legal_headers_splitter(text: str) -> List[str]:
    """
    Special splitter that preserves legal document structure.
    Fixed regex pattern with proper parenthesis balancing.
    """
    # Corrected regex pattern (added missing parenthesis)
    pattern = r'\n+(?:SECTION|Article|Clause)\s+[IVXLCDM0-9]+\.?\s*\n+'
    sections = re.split(pattern, text)

    return [s for s in sections if s.strip()]

def validate_file_extension(extension: str, allowed: set):
    if extension not in allowed:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Allowed: {', '.join(allowed)}"
        )